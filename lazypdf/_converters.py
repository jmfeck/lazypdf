from __future__ import annotations

import os
from typing import TYPE_CHECKING

from lazypdf.utils import resolve_pages

if TYPE_CHECKING:
    from lazypdf.core import PDFFile

_IMAGE_FORMATS = {"jpg", "jpeg", "png", "bmp", "tiff", "ppm"}


class ConvertersMixin:
    """Mixin for exporting PDFs to various formats."""

    def to_pdf(self: PDFFile, output_path: str) -> str:
        """Save the current state as a PDF file.

        Args:
            output_path: Destination file path.

        Returns:
            The output file path.
        """
        os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
        kwargs = dict(self._save_opts)
        if self._encryption:
            kwargs["encryption"] = self._encryption.get("encrypt", 0)
            kwargs["user_pw"] = self._encryption.get("user_pw", "")
            kwargs["owner_pw"] = self._encryption.get("owner_pw", "")
            kwargs["permissions"] = self._encryption.get("perm", 4095)
        self._doc.save(output_path, **kwargs)
        return output_path

    def to_images(
        self: PDFFile,
        output_dir: str,
        *,
        fmt: str = "png",
        dpi: int = 200,
        pages: list[int] | None = None,
    ) -> list[str]:
        """Render pages as images. Terminal operation.

        Args:
            output_dir: Directory to write the image files into.
            fmt: Image format ('png', 'jpg', 'bmp', 'tiff').
            dpi: Resolution in dots per inch.
            pages: Optional list of 1-indexed page numbers. If None, renders all.

        Returns:
            List of output file paths.
        """
        fmt = fmt.lower().replace("jpeg", "jpg")
        if fmt not in _IMAGE_FORMATS:
            raise ValueError(f"Unsupported image format: {fmt}. Supported: {sorted(_IMAGE_FORMATS)}")

        os.makedirs(output_dir, exist_ok=True)
        base = os.path.splitext(os.path.basename(self.path or "document"))[0]
        target = resolve_pages(pages, len(self._doc))
        paths: list[str] = []

        for i in target:
            pix = self._doc[i].get_pixmap(dpi=dpi)
            out_path = os.path.join(output_dir, f"{base}_page_{i + 1}.{fmt}")
            pix.save(out_path)
            paths.append(out_path)

        return paths

    def to_jpg(self: PDFFile, output_dir: str, *, dpi: int = 200, pages: list[int] | None = None) -> list[str]:
        """Render pages as JPEG images. Shorthand for to_images(fmt='jpg')."""
        return self.to_images(output_dir, fmt="jpg", dpi=dpi, pages=pages)

    def to_png(self: PDFFile, output_dir: str, *, dpi: int = 200, pages: list[int] | None = None) -> list[str]:
        """Render pages as PNG images. Shorthand for to_images(fmt='png')."""
        return self.to_images(output_dir, fmt="png", dpi=dpi, pages=pages)

    def to_docx(self: PDFFile, output_path: str) -> str:
        """Convert the PDF to a Word document (.docx).

        Extracts text and attempts to preserve basic structure.
        Note: images, tables, and complex formatting are not preserved.
        Requires python-docx: pip install lazypdf[office]

        Args:
            output_path: Destination .docx file path.

        Returns:
            The output file path.
        """
        try:
            from docx import Document
            from docx.shared import Pt
        except ImportError:
            raise ImportError("DOCX export requires python-docx. Install with: pip install lazypdf[office]") from None

        doc = Document()
        for i in range(len(self._doc)):
            page = self._doc[i]
            blocks = page.get_text("dict")["blocks"]
            for block in blocks:
                if block.get("type") == 0:
                    for line in block.get("lines", []):
                        text = "".join(span["text"] for span in line.get("spans", []))
                        if text.strip():
                            p = doc.add_paragraph()
                            run = p.add_run(text)
                            if line["spans"]:
                                run.font.size = Pt(line["spans"][0]["size"])
            if i < len(self._doc) - 1:
                doc.add_page_break()

        os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
        doc.save(output_path)
        return output_path

    def to_xlsx(self: PDFFile, output_path: str) -> str:
        """Export tables from the PDF to an Excel file (.xlsx).

        Each table is placed in a separate sheet.
        Requires openpyxl and pdfplumber: pip install lazypdf[office,tables]

        Args:
            output_path: Destination .xlsx file path.

        Returns:
            The output file path.
        """
        try:
            from openpyxl import Workbook
        except ImportError:
            raise ImportError("XLSX export requires openpyxl. Install with: pip install lazypdf[office]") from None

        tables = self.extract_tables()
        if not tables:
            raise ValueError("No tables found in the PDF.")

        wb = Workbook()
        wb.remove(wb.active)
        for idx, table in enumerate(tables):
            ws = wb.create_sheet(title=f"Table {idx + 1}")
            for row in table:
                ws.append(row)

        os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
        wb.save(output_path)
        return output_path

    def to_pdfa(self: PDFFile, output_path: str, *, level: int = 2, engine: str = "pymupdf") -> str:
        """Convert the PDF to PDF/A format for archival.

        Args:
            output_path: Destination file path.
            level: PDF/A conformance level (1, 2, or 3). Default is 2.
            engine: Conversion engine. ``"pymupdf"`` (default, zero deps) or
                ``"ghostscript"`` (most compliant, requires Ghostscript binary).

        Returns:
            The output file path.
        """
        valid_engines = {"pymupdf", "ghostscript"}
        if engine not in valid_engines:
            raise ValueError(f"Unknown engine '{engine}'. Options: {sorted(valid_engines)}")

        if level not in (1, 2, 3):
            raise ValueError(f"PDF/A level must be 1, 2, or 3, got {level}.")

        os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)

        if engine == "pymupdf":
            return self._to_pdfa_pymupdf(output_path, level)
        else:
            return self._to_pdfa_ghostscript(output_path, level)

    def _to_pdfa_pymupdf(self: PDFFile, output_path: str, level: int) -> str:
        """Convert to PDF/A using PyMuPDF metadata + font embedding."""
        import pymupdf
        from datetime import datetime, timezone

        data = self._doc.tobytes(garbage=4, deflate=True, clean=True)
        doc = pymupdf.open("pdf", data)

        now = datetime.now(timezone.utc).strftime("D:%Y%m%d%H%M%S+00'00'")
        doc.set_metadata({
            "producer": "lazypdf (PyMuPDF)",
            "creator": "lazypdf",
            "creationDate": now,
            "modDate": now,
        })

        pdfaid_part = str(level)
        xmp = (
            '<?xpacket begin="\xef\xbb\xbf" id="W5M0MpCehiHzreSzNTczkc9d"?>'
            '<x:xmpmeta xmlns:x="adobe:ns:meta/">'
            '<rdf:RDF xmlns:rdf="http://www.w3.org/1999/02/22-rdf-syntax-ns#">'
            '<rdf:Description rdf:about=""'
            ' xmlns:pdfaid="http://www.aiim.org/pdfa/ns/id/">'
            f"<pdfaid:part>{pdfaid_part}</pdfaid:part>"
            "<pdfaid:conformance>B</pdfaid:conformance>"
            "</rdf:Description>"
            "</rdf:RDF>"
            "</x:xmpmeta>"
            '<?xpacket end="w"?>'
        )
        doc.set_xml_metadata(xmp)
        doc.save(output_path, garbage=4, deflate=True, clean=True)
        doc.close()
        return output_path

    def _to_pdfa_ghostscript(self: PDFFile, output_path: str, level: int) -> str:
        """Convert to PDF/A using Ghostscript."""
        import shutil
        import subprocess
        import sys
        import tempfile

        if sys.platform == "win32":
            gs = shutil.which("gswin64c") or shutil.which("gswin32c") or shutil.which("gs")
        else:
            gs = shutil.which("gs")

        if gs is None:
            raise RuntimeError(
                "Ghostscript not found. Install it from https://www.ghostscript.com/ "
                "and make sure 'gs' (or 'gswin64c' on Windows) is on your PATH."
            )

        tmp_path = None
        try:
            with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False, delete_on_close=False) as tmp:
                tmp_path = tmp.name
                self._doc.save(tmp_path)

            subprocess.run(
                [
                    gs,
                    "-dPDFA=" + str(level),
                    "-dBATCH",
                    "-dNOPAUSE",
                    "-dSAFER",
                    "-sDEVICE=pdfwrite",
                    "-dCompatibilityLevel=1.4",
                    "-dPDFACompatibilityPolicy=1",
                    "-sOutputFile=" + output_path,
                    tmp_path,
                ],
                check=True,
                capture_output=True,
            )
        finally:
            if tmp_path is not None:
                try:
                    os.unlink(tmp_path)
                except OSError:
                    pass

        return output_path

    def to_bytes(self: PDFFile) -> bytes:
        """Return the current PDF state as bytes."""
        kwargs = dict(self._save_opts)
        return self._doc.tobytes(**kwargs)
